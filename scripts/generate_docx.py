"""
generate_docx.py  -  CENG 479 Sub2 Final Report (Word .docx)
Usage: python scripts/generate_docx.py
Output: CENG479_Sub2_Final_Report.docx
"""

import csv, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GITHUB = "https://github.com/Muhammedcakirgoz/parallel-image-processing"

# Colors (RGB)
C_BLUE      = RGBColor(0x1a, 0x3a, 0x6b)
C_BLUE_MID  = RGBColor(0x2e, 0x6d, 0xb4)
C_GOLD      = RGBColor(0xf0, 0xc0, 0x40)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY      = RGBColor(0x44, 0x44, 0x44)
C_GREEN     = RGBColor(0x1a, 0x6b, 0x3a)
C_ACCENT    = RGBColor(0xd6, 0xe4, 0xf7)
C_BLACK     = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(table, color_hex="C8D8F0"):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color_hex)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def paragraph_shade(para, hex_color):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def set_para_border_bottom(para, color_hex, sz="12"):
    pPr     = para._p.get_or_add_pPr()
    pBdr    = OxmlElement("w:pBdr")
    bottom  = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl   = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r    = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E6DB4")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------
def load_speedup():
    path = os.path.join(BASE, "speedup_table.csv")
    rows = []
    with open(path, newline="", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            rows.append(r)
    return rows


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------
def set_run_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=C_BLUE, size=14, border=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    set_run_font(run, size, bold=True, color=color)
    if border:
        set_para_border_bottom(para, "1a3a6b", "12" if level == 1 else "6")
    return para


def add_body(doc, text, size=9.5, justify=True, color=C_BLACK, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run_font(run, size, color=color)
    return para


def add_body_mixed(doc, parts, size=9.5, justify=True, space_after=4):
    """parts: list of (text, bold, italic, color, url)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic, color, url in parts:
        if url:
            add_hyperlink(para, text, url)
        else:
            run = para.add_run(text)
            set_run_font(run, size, bold=bold, italic=italic, color=color or C_BLACK)
    return para


def add_bullet(doc, text, size=9.5):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.left_indent  = Cm(0.8)
    run = para.add_run(text)
    set_run_font(run, size)
    return para


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.space_before = Pt(2)
    run = para.add_run(text)
    set_run_font(run, 8, italic=True, color=C_GRAY)
    return para


def add_spacer(doc, pt=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(pt)
    para.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------
def add_data_table(doc, headers, data_rows, col_widths_cm, hdr_color="1a3a6b", hdr_text_color=C_WHITE):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, hdr_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        set_run_font(run, 8.5, bold=True, color=hdr_text_color)

    # data rows
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        bg  = "FFFFFF" if ri % 2 == 0 else "EAF1FB"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            set_run_font(run, 8)

    set_cell_borders(table)
    return table


# ---------------------------------------------------------------------------
# Slide box (shaded paragraph block)
# ---------------------------------------------------------------------------
def add_slide(doc, number, title, items, note=None):
    # slide header
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(0)
    paragraph_shade(para, "1a3a6b")
    badge = para.add_run(f"  {number}  ")
    set_run_font(badge, 9, bold=True, color=C_GOLD)
    title_run = para.add_run(f"  {title}")
    set_run_font(title_run, 12, bold=True, color=C_WHITE)

    # items
    for kind, text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        paragraph_shade(p, "1e4070")
        if kind == "bullet":
            p.paragraph_format.left_indent = Cm(0.8)
            dot = p.add_run("  •  ")
            set_run_font(dot, 9, color=C_GOLD)
            run = p.add_run(text)
            set_run_font(run, 9, color=C_WHITE)
        elif kind == "sub":
            p.paragraph_format.left_indent = Cm(0.4)
            run = p.add_run(f"  {text}")
            set_run_font(run, 9, bold=True, color=C_GOLD)
        elif kind == "body":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, 9, italic=True, color=RGBColor(0xaa, 0xcc, 0xee))
        elif kind == "space":
            p.paragraph_format.space_after = Pt(3)

    # note
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        paragraph_shade(p, "0d2244")
        run = p.add_run(f"  {note}  ")
        set_run_font(run, 7.5, italic=True, color=C_GOLD)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)

    # default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    return doc


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def build_cover(doc):
    # top colour block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "F0C040")
    r = p.add_run("  CENG 479 — Parallel Computing  |  Gazi University  |  Spring 2026  ")
    set_run_font(r, 10, bold=True, color=C_BLUE)

    # main title block
    for text, sz, bold, clr, bg in [
        ("", 4, False, C_WHITE, "1a3a6b"),
        ("Parallel Image Processing Engine", 22, True,  C_WHITE,  "1a3a6b"),
        ("Final Implementation Report & Presentation", 13, False, C_GOLD, "1a3a6b"),
        ("", 4, False, C_WHITE, "1a3a6b"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, bg)
        if text:
            r = p.add_run(text)
            set_run_font(r, sz, bold=bold, color=clr)

    # divider
    p = doc.add_paragraph()
    paragraph_shade(p, "1a3a6b")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    set_para_border_bottom(p, "F0C040", "12")

    # badges row
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    paragraph_shade(p, "1a3a6b")
    for badge in ["Java 17", "Maven 3.6+", "JMH", "ExecutorService", "ForkJoinPool"]:
        r = p.add_run(f"  {badge}  ")
        set_run_font(r, 9, bold=True, color=C_WHITE)

    # filters
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "1a3a6b")
    r = p.add_run("Filters Implemented")
    set_run_font(r, 10, bold=True, color=C_GOLD)

    for fname, fdesc in [
        ("Grayscale",         "Point-wise — Low compute"),
        ("Gaussian Blur 5x5", "Convolution — High compute (25 taps/pixel)"),
        ("Sobel 3x3",         "Gradient — Edge detection"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, "1a3a6b")
        dot = p.add_run("● ")
        set_run_font(dot, 9, color=C_GOLD)
        rb = p.add_run(fname)
        set_run_font(rb, 9, bold=True, color=C_WHITE)
        rd = p.add_run(f"   {fdesc}")
        set_run_font(rd, 9, color=RGBColor(0xaa, 0xcc, 0xee))

    # team box
    for text, sz, bold, clr, bg in [
        ("", 3, False, C_WHITE, "0d2244"),
        ("Team Members", 10, True,  C_GOLD,  "0d2244"),
        ("Muhammed Çakırgöz   ·   Musa Bilal Yaz", 13, True, C_WHITE, "0d2244"),
        ("Gazi University  —  Computer Engineering  —  Submission 2", 9, False, RGBColor(0xaa,0xcc,0xee), "0d2244"),
        ("", 3, False, C_WHITE, "0d2244"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, bg)
        if text:
            r = p.add_run(text)
            set_run_font(r, sz, bold=bold, color=clr)

    add_spacer(doc, 4)

    # github box
    for text, sz, bold, clr, bg, url in [
        ("", 3, False, C_WHITE, "0d2244", None),
        ("GitHub Repository (Public)", 10, True, C_GOLD, "0d2244", None),
        (GITHUB, 9, False, RGBColor(0xaa,0xcc,0xee), "0d2244", GITHUB),
        ("Source code, benchmarks, and all report assets are available at the link above",
         8, False, RGBColor(0xaa,0xcc,0xee), "0d2244", None),
        ("", 3, False, C_WHITE, "0d2244", None),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, bg)
        if text and not url:
            r = p.add_run(text)
            set_run_font(r, sz, bold=bold, color=clr)
        elif url:
            add_hyperlink(p, text, url)

    # bottom bar
    add_spacer(doc, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "F0C040")
    r = p.add_run("  June 2026  —  CENG 479 Parallel Computing  ")
    set_run_font(r, 10, bold=True, color=C_BLUE)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------
def build_report(doc, rows):
    # page title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("CENG 479 — Parallel Image Processing Engine")
    set_run_font(r, 16, bold=True, color=C_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("Implementation Report  |  Submission 2")
    set_run_font(r2, 9, italic=True, color=C_GRAY)
    set_para_border_bottom(p2, "1a3a6b", "6")

    # 1. Introduction
    add_heading(doc, "1. Introduction")
    add_body(doc,
        "High-resolution image processing is computationally demanding. Sequential execution of "
        "kernel-based convolution filters creates significant bottlenecks at large image sizes (4K+). "
        "Because each output pixel depends solely on a fixed neighbourhood of the input image, "
        "the problem is embarrassingly parallel: the image can be divided across threads with "
        "zero inter-thread data dependencies during the compute phase.")
    add_body(doc,
        "This report presents a parallel engine built in Java using two concurrency strategies: "
        "ExecutorService with horizontal strip decomposition and ForkJoinPool with "
        "divide-and-conquer work-stealing. Three filters — Grayscale, Gaussian Blur 5x5, Sobel 3x3 — "
        "were benchmarked with JMH across image sizes 512x512, 1024x1024, 2048x2048.")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Source code (public): ")
    set_run_font(r, 9.5)
    add_hyperlink(p, GITHUB, GITHUB)

    # 2. Sequential Baseline
    add_heading(doc, "2. Sequential Baseline Implementation")
    add_body(doc,
        "SequentialProcessor iterates row-by-row applying the given filter. "
        "The Filter interface exposes a single apply(pixels[], width, height, x, y) method. "
        "Edge handling uses clamped coordinates, ensuring identical boundary behaviour "
        "across all implementations.")
    add_spacer(doc, 3)
    add_data_table(doc,
        ["Filter", "Type", "Kernel", "Cost per pixel"],
        [
            ["GrayscaleFilter",   "Point-wise",  "None (r=0)", "Low — 3 reads, 1 write"],
            ["GaussianBlurFilter","Convolution",  "5x5",        "High — 25 weighted taps"],
            ["SobelFilter",       "Gradient",     "3x3 Gx+Gy",  "Medium — 18 taps + sqrt"],
        ],
        [4.5, 3.0, 2.5, 5.5])
    add_spacer(doc, 6)

    # 3. Parallel Implementation
    add_heading(doc, "3. Parallel Implementation")
    add_heading(doc, "3.1  ExecutorService — Horizontal Strip Decomposition", level=2, color=C_BLUE_MID, size=11, border=False)
    add_body(doc,
        "ExecutorParallelProcessor divides the image into N equal horizontal strips "
        "(one per thread). Each Callable task processes rows [start, end) and writes into a "
        "pre-allocated output array at its assigned offset. Strips are disjoint and the source "
        "array is read-only, so no locking is required. The thread pool is created once and "
        "reused via try-with-resources.")
    add_heading(doc, "3.2  ForkJoinPool — Divide-and-Conquer", level=2, color=C_BLUE_MID, size=11, border=False)
    add_body(doc,
        "ForkJoinParallelProcessor uses a RecursiveAction that halves the row range until it "
        "falls below 64 rows (sequential threshold), then processes directly. "
        "The work-stealing scheduler rebalances load dynamically — advantageous for larger "
        "images where static strip imbalance would arise.")
    add_heading(doc, "3.3  Correctness Verification", level=2, color=C_BLUE_MID, size=11, border=False)
    add_body(doc,
        "CorrectnessVerifier.firstDifference() performs a pixel-for-pixel comparison between "
        "parallel output and sequential baseline. All 6 combinations (3 filters x 2 strategies) "
        "on a 2048x2048 synthetic image produce bit-identical output:")
    add_spacer(doc, 3)
    add_data_table(doc,
        ["Filter", "ExecutorService (12 threads)", "ForkJoinPool (12 threads)"],
        [
            ["Grayscale",       "PASS", "PASS"],
            ["GaussianBlur5x5", "PASS", "PASS"],
            ["Sobel3x3",        "PASS", "PASS"],
        ],
        [5.0, 5.5, 5.0], hdr_color="1a6b3a")
    add_spacer(doc, 6)

    # 4. Performance
    add_heading(doc, "4. Performance Comparison")
    add_body(doc,
        "All benchmarks use JMH (10 warm-up + 10 measurement iterations, AverageTime mode, ms/op) "
        "on a 12-logical-core machine. Speedup = T_sequential / T_parallel.")

    for fname, label in [
        ("GaussianBlur5x5", "4.1  GaussianBlur 5x5"),
        ("Sobel3x3",        "4.2  Sobel 3x3"),
        ("Grayscale",       "4.3  Grayscale"),
    ]:
        add_heading(doc, label, level=2, color=C_BLUE_MID, size=11, border=False)
        data = []
        for r in rows:
            if r["filter"] == fname:
                data.append([
                    f'{r["size"]}^2', r["threads"],
                    r["sequential_ms"], r["executor_ms"], r["forkjoin_ms"],
                    r["executor_speedup"]+"x", r["forkjoin_speedup"]+"x",
                ])
        add_data_table(doc,
            ["Size","Threads","Seq (ms)","Executor (ms)","ForkJoin (ms)","Exec x","FJ x"],
            data,
            [2.0, 1.7, 2.4, 2.6, 2.6, 1.9, 1.8])
        add_spacer(doc, 4)

    add_heading(doc, "4.4  Speedup Charts", level=2, color=C_BLUE_MID, size=11, border=False)
    add_spacer(doc, 2)

    chart_w = Inches(6.0)
    for fname, caption in [
        ("speedup_GaussianBlur5x5.png",
         "Figure 1 — GaussianBlur 5x5: ExecutorService (left) vs ForkJoinPool (right)"),
        ("speedup_Sobel3x3.png",
         "Figure 2 — Sobel 3x3: ExecutorService (left) vs ForkJoinPool (right)"),
        ("speedup_Grayscale.png",
         "Figure 3 — Grayscale: memory-bound, speedup plateaus ~1.5x"),
        ("speedup_combined.png",
         "Figure 4 — All filters combined (2048x2048, Executor)"),
    ]:
        path = os.path.join(BASE, fname)
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.add_run().add_picture(path, width=chart_w)
        add_caption(doc, caption)

    add_heading(doc, "4.5  Analysis", level=2, color=C_BLUE_MID, size=11, border=False)
    add_body(doc,
        "Compute-bound filters (Gaussian Blur, Sobel) scale well, reaching 4.3x–5.1x at 8 threads. "
        "ForkJoinPool outperforms ExecutorService on larger images due to dynamic work-stealing load balancing.")
    add_body(doc,
        "Memory-bound filter (Grayscale) plateaus at ~1.5x regardless of thread count. "
        "The single-channel lookup saturates memory bandwidth before CPU utilisation can increase — "
        "consistent with the Roofline Model: low arithmetic intensity = bandwidth-limited.")
    add_body(doc,
        "Amdahl's Law: the empirical ceiling for Gaussian Blur aligns with p ~ 0.95 parallel "
        "fraction; sequential overhead (array allocation, thread setup) limits the theoretical max.")

    # 5. Academic Background
    add_heading(doc, "5. Academic Background")
    add_body(doc,
        "Seinstra et al. (2002) showed convolution filters exhibit near-linear speedup when "
        "partitioning minimises inter-thread communication — consistent with our strip-decomposition "
        "results for compute-intensive kernels.")
    add_body(doc,
        "Amdahl (1967): S = 1/(1-p). For Gaussian Blur, p ~ 0.95 gives theoretical max ~20x; "
        "our 4-5x at 8 threads is expected given practical overhead.")
    add_body(doc,
        "Williams et al. (2009) — Roofline Model — explains divergent scalability via arithmetic "
        "intensity: memory-bound kernels saturate bandwidth before compute capacity. "
        "Lea (2000) provides the theoretical basis for ForkJoinPool's work-stealing scheduler.")

    # 6. Challenges
    add_heading(doc, "6. Challenges and Solutions")
    for title, desc in [
        ("JVM JIT Warm-up and GC Noise",
         "System.nanoTime() measurements were skewed by JIT and GC. "
         "Solution: Adopted JMH with warm-up iterations and fork-per-config."),
        ("Bit-Identical Correctness",
         "Strip boundaries required careful handling. "
         "Solution: CorrectnessVerifier with pixel-for-pixel scan; "
         "clamped-coordinate edge handling ensures identical border treatment."),
        ("Varying Scalability Between Filters",
         "Grayscale achieved only ~1.5x with 8 threads. "
         "Solution: Profiling confirmed memory bandwidth saturation (Roofline Model) "
         "as root cause — not thread management inefficiency."),
        ("ForkJoin Recursive Overhead on Small Images",
         "For 512x512, ForkJoin recursion overhead slightly exceeded Executor. "
         "Solution: Tuned SEQUENTIAL_THRESHOLD to 64 rows."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        rb = p.add_run(f"• {title}: ")
        set_run_font(rb, 9.5, bold=True)
        rd = p.add_run(desc)
        set_run_font(rd, 9.5)

    # 7. Conclusion
    add_heading(doc, "7. Conclusion and Future Improvements")
    add_body(doc,
        "This project successfully demonstrated parallelisation of image convolution algorithms "
        "using Java concurrency. JMH benchmarks confirmed significant speedups for compute-bound "
        "operations: Gaussian Blur reached ~4.30x (Executor) and ~4.81x (ForkJoin) at 8 threads "
        "on 2048x2048; Sobel reached ~4.49x and ~5.10x. Grayscale was confirmed memory-bound "
        "at ~1.55x. ForkJoinPool's work-stealing advantage grew with image size.")
    for fw in [
        "GPU/CUDA — exploit thousands of CUDA cores for pixel-independent kernels",
        "SIMD vectorisation (AVX2/AVX-512) — intra-core throughput gains",
        "4K and 8K benchmarks — evaluate boundary overhead at ultra-high resolution",
        "Adaptive thread pool — dynamic tuning based on image size and hardware profile",
        "Cache-blocking and prefetching — optimise memory-bound filters like Grayscale",
    ]:
        add_bullet(doc, fw)

    # 8. References
    add_heading(doc, "8. References")
    refs = [
        "Amdahl, G. M. (1967). Validity of the single processor approach to achieving large scale "
        "computing capabilities. Proc. Spring Joint Computer Conf., 483-485. "
        "https://doi.org/10.1145/1465482.1465560",
        "Seinstra, F. J., Koelma, D., & Bagdanov, A. D. (2002). Finite state machine-based "
        "optimization of data parallel regular domain problems applied to low-level image "
        "processing. IEEE Trans. Parallel Distrib. Syst., 15(10), 865-877. "
        "https://doi.org/10.1109/TPDS.2004.45",
        "Williams, S., Waterman, A., & Patterson, D. (2009). Roofline: An insightful visual "
        "performance model for multicore architectures. Commun. ACM, 52(4), 65-76. "
        "https://doi.org/10.1145/1498765.1498785",
        "Lea, D. (2000). A Java fork/join framework. Proc. ACM Java Grande Conf., 36-43. "
        "https://doi.org/10.1145/337449.337465",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        r = p.add_run(f"[{i}]  {ref}")
        set_run_font(r, 8.5, color=C_GRAY)


# ---------------------------------------------------------------------------
# Helper: write slide content into a table cell
# ---------------------------------------------------------------------------
def fill_slide_cell(cell, number, title, items, note):
    """Write a compact slide into a Word table cell."""

    def cp(text, size, bold=False, italic=False, color=C_WHITE,
           bg="1a3a6b", align=WD_ALIGN_PARAGRAPH.LEFT, indent=0):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        paragraph_shade(p, bg)
        if text:
            r = p.add_run(text)
            set_run_font(r, size, bold=bold, italic=italic, color=color)
        return p

    # clear default empty paragraph that Word adds
    for para in list(cell.paragraphs):
        p = para._p
        p.getparent().remove(p)

    # title bar
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "0d2244")
    badge = p.add_run(f" {number} ")
    set_run_font(badge, 7, bold=True, color=C_GOLD)
    t = p.add_run(f"  {title}")
    set_run_font(t, 8, bold=True, color=C_WHITE)

    # items
    for kind, text in items:
        if kind == "bullet":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.4)
            paragraph_shade(p, "1a3a6b")
            dot = p.add_run("• ")
            set_run_font(dot, 7.5, color=C_GOLD)
            r = p.add_run(text)
            set_run_font(r, 7.5, color=C_WHITE)
        elif kind == "sub":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.2)
            paragraph_shade(p, "1a3a6b")
            r = p.add_run(text)
            set_run_font(r, 7.5, bold=True, color=C_GOLD)
        elif kind == "body":
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            paragraph_shade(p, "1a3a6b")
            r = p.add_run(text)
            set_run_font(r, 7, italic=True, color=RGBColor(0xaa, 0xcc, 0xee))
        elif kind == "space":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            paragraph_shade(p, "1a3a6b")

    # note bar
    if note:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, "0d2244")
        r = p.add_run(f" {note} ")
        set_run_font(r, 6.5, italic=True, color=C_GOLD)


# ---------------------------------------------------------------------------
# Full-page slide content
# ---------------------------------------------------------------------------
def _fill_slide_full(cell, number, title, items, note):
    """Write a full-page slide into a Word table cell."""

    for para in list(cell.paragraphs):
        para._p.getparent().remove(para._p)

    def cp(text, size, bold=False, italic=False, color=C_WHITE,
           bg="1a3a6b", align=WD_ALIGN_PARAGRAPH.LEFT, indent=0, space_after=0):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        paragraph_shade(p, bg)
        if text:
            r = p.add_run(text)
            set_run_font(r, size, bold=bold, italic=italic, color=color)
        return p

    # top gold bar
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "F0C040")
    r = p.add_run(f"  CENG 479 — Parallel Computing  |  Slide {number} / 10  ")
    set_run_font(r, 9, bold=True, color=C_BLUE)

    # spacer
    cp("", 6, bg="1a3a6b")

    # slide number badge + title
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    paragraph_shade(p, "0d2244")
    badge = p.add_run(f"  {number}  ")
    set_run_font(badge, 14, bold=True, color=C_GOLD)
    t = p.add_run(f"   {title}")
    set_run_font(t, 16, bold=True, color=C_WHITE)

    # gold divider line
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    paragraph_shade(p, "1a3a6b")
    set_para_border_bottom(p, "F0C040", "12")

    # items
    for kind, text in items:
        if kind == "bullet":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(1.2)
            paragraph_shade(p, "1a3a6b")
            dot = p.add_run("●  ")
            set_run_font(dot, 11, color=C_GOLD)
            r = p.add_run(text)
            set_run_font(r, 11, color=C_WHITE)
        elif kind == "sub":
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.6)
            paragraph_shade(p, "1a3a6b")
            r = p.add_run(text)
            set_run_font(r, 11, bold=True, color=C_GOLD)
        elif kind == "body":
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            paragraph_shade(p, "1a3a6b")
            r = p.add_run(text)
            set_run_font(r, 10, italic=True, color=RGBColor(0xaa, 0xcc, 0xee))
        elif kind == "space":
            cp("", 8, bg="1a3a6b", space_after=4)

    # push note to bottom with spacer
    cp("", 6, bg="1a3a6b", space_after=8)

    # note bar at bottom
    if note:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        paragraph_shade(p, "0d2244")
        r = p.add_run(f"  {note}  ")
        set_run_font(r, 9, italic=True, color=C_GOLD)

    # bottom gold bar
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    paragraph_shade(p, "F0C040")
    r = p.add_run(f"  June 2026  —  Gazi University  |  Computer Engineering  ")
    set_run_font(r, 9, bold=True, color=C_BLUE)


# ---------------------------------------------------------------------------
# Slides  —  each slide on its own page
# ---------------------------------------------------------------------------
def build_slides(doc):
    doc.add_page_break()
    add_heading(doc, "PRESENTATION SLIDES  (10 Slides)", size=13)
    add_spacer(doc, 2)

    slides = [
        (1, "Parallel Image Processing Engine", [
            ("space",""),
            ("body",  "CENG 479 — Parallel Computing  |  Gazi University  |  Spring 2026"),
            ("body",  "Muhammed Cakiagoz   ·   Musa Bilal Yaz"),
            ("space",""),
            ("body",  GITHUB),
        ], "Java Threads | ExecutorService | ForkJoinPool | JMH Benchmark"),

        (2, "Problem & Motivation", [
            ("bullet","High-resolution image filtering (4K+) is slow on a single thread"),
            ("bullet","Convolution cost: O(W x H x K^2) per filter — bottleneck at scale"),
            ("bullet","Each output pixel is INDEPENDENT — embarrassingly parallel"),
            ("bullet","Goal: multi-threaded speedup with zero correctness compromise"),
            ("space",""),
            ("body",  "Amdahl's Law: S = 1/(1-p)  —  p ~ 0.95 for Gaussian Blur"),
        ], "Theoretical max speedup ~20x at p=0.95"),

        (3, "Architecture Overview", [
            ("sub",   "Filters:"),
            ("bullet","Grayscale (point-wise), Gaussian Blur 5x5, Sobel 3x3"),
            ("sub",   "Implementations:"),
            ("bullet","SequentialProcessor — single-thread baseline"),
            ("bullet","ExecutorParallelProcessor — fixed pool + horizontal strips"),
            ("bullet","ForkJoinParallelProcessor — RecursiveAction + work-stealing"),
            ("sub",   "Tools:"),
            ("bullet","CorrectnessVerifier + JMH Benchmark"),
        ], "All source code publicly available on GitHub"),

        (4, "ExecutorService Design", [
            ("bullet","Image split into N equal horizontal strips (N = thread count)"),
            ("bullet","Each Callable processes rows [start, end) with no locking"),
            ("bullet","Source array is READ-ONLY — zero synchronisation overhead"),
            ("bullet","Thread pool created once, reused via try-with-resources"),
            ("space",""),
            ("body",  "Strip[i] covers rows:  i*(H/N)  to  (i+1)*(H/N)"),
        ], "Static decomposition — optimal for uniform workloads"),

        (5, "ForkJoinPool Design", [
            ("bullet","RecursiveAction splits row range in half recursively"),
            ("bullet","Base case: range <= 64 rows — process sequentially"),
            ("bullet","Work-stealing: idle threads steal tasks from busy queues"),
            ("bullet","Better dynamic load balance than static strips for large images"),
            ("space",""),
            ("bullet","ForkJoin outperforms Executor on 2048x2048 (Gaussian: 4.81x vs 4.30x)"),
        ], "Dynamic decomposition — adapts to runtime load imbalance"),

        (6, "Correctness Verification", [
            ("body",  "CorrectnessVerifier.firstDifference() — pixel-by-pixel scan"),
            ("space",""),
            ("body",  "Tested on 2048x2048 synthetic image — ALL 6 combinations PASS"),
            ("space",""),
            ("bullet","Grayscale:       Executor PASS  |  ForkJoin PASS"),
            ("bullet","GaussianBlur5x5: Executor PASS  |  ForkJoin PASS"),
            ("bullet","Sobel3x3:        Executor PASS  |  ForkJoin PASS"),
        ], "Bit-identical output to sequential baseline in all cases"),

        (7, "Performance Results (8 threads, 2048x2048)", [
            ("sub",   "Compute-bound (scales well):"),
            ("bullet","Gaussian Blur:  Executor 4.30x  |  ForkJoin 4.81x"),
            ("bullet","Sobel 3x3:      Executor 4.49x  |  ForkJoin 5.10x"),
            ("space",""),
            ("sub",   "Memory-bound (bandwidth limited):"),
            ("bullet","Grayscale:      Executor 1.55x  |  ForkJoin 1.54x"),
        ], "JMH | 10 warm-up + 10 measurement iterations | AverageTime ms/op"),

        (8, "Speedup Charts", [
            ("space",""),
            ("body",  "GaussianBlur5x5 & Sobel — near-linear scaling up to 4-5x at 8 threads"),
            ("body",  "Grayscale — memory bandwidth saturation caps speedup at ~1.5x"),
            ("body",  "ForkJoin advantage grows with image size (work-stealing benefit)"),
            ("space",""),
            ("body",  "[See Figures 1-4 in the report section for full charts]"),
        ], "Roofline Model: compute-bound = scales | memory-bound = does not"),

        (9, "Challenges & Key Findings", [
            ("bullet","JIT warm-up noise — solved with JMH (fork-per-config, 10 warm-up iters)"),
            ("bullet","Correctness — clamped edges + pixel-for-pixel CorrectnessVerifier"),
            ("bullet","Grayscale bottleneck — memory bandwidth saturated (Roofline Model)"),
            ("bullet","ForkJoin overhead on small images — tuned threshold to 64 rows"),
            ("space",""),
            ("body",  "Key insight: speedup is FILTER-DEPENDENT — arithmetic intensity determines scalability"),
        ], "Benchmark: 12 logical cores, JDK 17, Maven 3.6"),

        (10, "Conclusion & Future Work", [
            ("bullet","Parallelised 3 convolution filters with 2 Java concurrency strategies"),
            ("bullet","Up to 5.10x speedup (Sobel, ForkJoin, 8 threads, 2048x2048)"),
            ("bullet","Memory-bound filters need different strategy (not thread-level parallelism)"),
            ("sub",   "Future Work:"),
            ("bullet","GPU/CUDA — thousands of concurrent cores for pixel-independent kernels"),
            ("bullet","SIMD vectorisation (AVX2/AVX-512) for intra-core throughput"),
            ("bullet","Adaptive thread pool — dynamic tuning per image size and hardware"),
        ], GITHUB),
    ]

    # Her slayt ayrı sayfada, tam sayfa tablo
    for number, title, items, note in slides:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        cell.width = Cm(16.0)
        set_cell_bg(cell, "1a3a6b")
        _fill_slide_full(cell, number, title, items, note)
        set_cell_borders(tbl, color_hex="2e6db4")
        if number < len(slides):
            doc.add_page_break()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    rows = load_speedup()
    doc  = setup_doc()

    build_cover(doc)
    build_report(doc, rows)
    build_slides(doc)

    out = os.path.join(BASE, "CENG479_Sub2_Final_Report.docx")
    doc.save(out)
    size_kb = os.path.getsize(out) // 1024
    print(f"DOCX written: {out}  ({size_kb} KB)")


if __name__ == "__main__":
    main()
